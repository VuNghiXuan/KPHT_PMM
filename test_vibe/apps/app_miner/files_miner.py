import re
import openpyxl
import datetime
import traceback
import logging
import json
import time
import pandas as pd 
from docx import Document 
from apps.app_knowledge.ai_gateway import AIGateway
from openpyxl.cell.cell import MergedCell
from django.db import transaction

# Import các Model đã cập nhật
from .models import DataEntry, DataField, ExcelTableRegion
from apps.app_knowledge.models import KnowledgeDraft
from apps.app_coach.models import DataType
# Import module phân tích vùng vừa tách file
from .excel_region_analyzer import ExcelRegionAnalyzer

logger = logging.getLogger(__name__)

class DataMinerService:
    def __init__(self):
        # Danh sách các từ khóa để nhận diện UI (Nút bấm, rác)
        self.ui_keywords = [
            'lưu', 'tìm kiếm', 'xóa', 'thêm', 'in phiếu', 'thoát', 'đóng', 
            'update', 'filter', 'search', 'edit', 'save', 'print'
        ]
        # Màu sắc quy ước cho các trường bắt buộc (Tiệm vàng thường dùng)
        self.required_colors = ["FFFF0000", "FFFFC000"] # Đỏ, Cam

    @classmethod
    def run_workflow(cls, project):
        miner = cls()
        project.status = 'PROCESSING'
        project.save()
        
        success, message = miner.process_project(project)
        
        project.status = 'COMPLETED' if success else 'FAILED'
        project.save()
        return success, message

    def process_project(self, project):
        try:
            # Điều phối dựa trên file_type của dự án
            if project.file_type == 'EXCEL':
                success, message = self._process_excel_file(project)
            elif project.file_type in ['TEXT', 'CSV']:
                success, message = self._process_txt_csv(project.file_path.path)
            elif project.file_type == 'DOCX':
                success, message = self._process_docx(project.file_path.path)
            elif project.file_type == 'IMAGE':
                success, message = self._process_image_easyocr(project.file_path.path, project)
            else:
                return False, f"Unsupported file type: {project.file_type}"

            return success, message

        except Exception as e:
            logger.error(f"Lỗi Miner tổng hợp: {str(e)}")
            return False, traceback.format_exc()

    def _process_excel_file(self, project):
        try:
            wb = openpyxl.load_workbook(project.file_path.path, data_only=False)
            total_sheets = len(wb.sheetnames)

            for index, sheet_name in enumerate(wb.sheetnames):
                ws = wb[sheet_name]
                print(f"--- [Miner] Đang bóc tách Sheet: {sheet_name} ({index+1}/{total_sheets}) ---")
                
                # 1. Khởi tạo/Cập nhật DataEntry
                sheet_obj, _ = DataEntry.objects.update_or_create(
                    project=project, 
                    name=sheet_name,
                    defaults={'refine_status': 'PENDING'} 
                )

                # Dọn dẹp dữ liệu ô cũ để tránh trùng lặp
                DataField.objects.filter(sheet=sheet_obj).delete()
                
                # 2. Container gom Metadata thô
                sheet_intel = {
                    "logic_blocks": [],   
                    "ui_elements": [],    
                    "business_data": [],  
                    "raw_structure": []   
                }

                fields_to_create = []
                merged_ranges = ws.merged_cells.ranges

                # 3. Duyệt toàn bộ các ô trong bản tính để gom hạt cát dữ liệu
                for row in ws.iter_rows(max_row=ws.max_row, max_col=ws.max_column):
                    if not any(cell.value is not None for cell in row):
                        continue 

                    for cell in row:
                        if any(cell.coordinate in r and cell.coordinate != r.start_cell.coordinate for r in merged_ranges):
                            continue

                        if cell.value is not None or (cell.comment):
                            field_type, label = self._classify_cell(ws, cell)
                            field_obj = self._build_field_obj(sheet_obj, cell, field_type, label)
                            fields_to_create.append(field_obj)

                            item_info = {
                                "address": cell.coordinate,
                                "label": label,
                                "value": str(cell.value) if not field_obj.formula else "FORMULA",
                                "formula": field_obj.formula,
                                "comment": cell.comment.text if cell.comment else None
                            }

                            if field_type == 'LOGIC':
                                sheet_intel["logic_blocks"].append(item_info)
                            elif field_type == 'UI':
                                sheet_intel["ui_elements"].append(item_info)
                            elif field_type == 'DATA':
                                sheet_intel["business_data"].append(item_info)

                # 4. Lưu Bulk DataFields vào DB trước để lấy dữ liệu nền tảng
                if fields_to_create:
                    DataField.objects.bulk_create(fields_to_create)

                # ⭐ [TÍNH NĂNG MỚI]: Triệu hồi module Analyzer vừa tách file để tự động gom Vùng Nghiệp Vụ
                ExcelRegionAnalyzer.cluster_and_bind_regions(sheet_obj)

                # 5. Đóng gói Metadata tổng hợp vào Sheet
                sheet_obj.metadata = sheet_intel
                sheet_obj.refine_status = 'EXTRACTED'
                
                sheet_obj.description = (
                    f"Sheet có {len(sheet_intel['logic_blocks'])} logic tính toán, "
                    f"{len(sheet_intel['business_data'])} trường nghiệp vụ. "
                    f"Đã nhận diện {len(sheet_intel['ui_elements'])} thành phần UI."
                )
                sheet_obj.save()

                # 6. Tạo duy nhất 1 bản thảo gom tri thức phục vụ cho RAG
                self._create_unified_draft(project, sheet_obj)

            return True, f"Thành công! Đã bóc tách cấu trúc và phân vùng cho {total_sheets} sheets."

        except Exception as e:
            logger.error(f"Lỗi xử lý file Excel: {str(e)}")
            return False, traceback.format_exc()

    def _process_txt_csv(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return True, {"status": "success", "data": content}
        except Exception as e:
            logger.error(f"Lỗi xử lý TXT/CSV: {str(e)}")
            return False, traceback.format_exc()

    def _process_docx(self, file_path):
        try:
            document = Document(file_path)
            full_text = []
            for para in document.paragraphs:
                full_text.append(para.text)
            for table in document.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    full_text.append("\t".join(row_text))
            content = "\n".join(full_text)
            return True, {"status": "success", "data": content}
        except Exception as e:
            logger.error(f"Lỗi xử lý DOCX: {str(e)}")
            return False, traceback.format_exc()

    def _process_image_easyocr(self, file_path, project):
        try:
            import easyocr
            print(f"--- [EasyOCR] Đang quét chữ trên hình ảnh: {project.name} ---")
            
            reader = easyocr.Reader(['vi', 'en'], gpu=False)
            results = reader.readtext(file_path, detail=0)
            raw_text = "\n".join(results)
            
            if not raw_text.strip():
                return False, "EasyOCR không tìm thấy hoặc không đọc được chữ nào từ hình ảnh."
            
            ai_gateway = AIGateway()
            system_prompt = (
                "Bạn là một trợ lý AI chuyên gia phân tích dữ liệu tiệm vàng (Ứng Dụng Vàng).\n"
                "Nhiệm vụ của bạn là đọc đoạn văn bản thô thu được từ máy quét OCR bên dưới, "
                "bóc tách thành cấu trúc JSON chuẩn nghiệp vụ bao gồm: tên_khách_hàng, ngày_giao_dịch, "
                "danh_sách_sản_phẩm (tên_món, loại_vàng, trọng_lượng, tiền_công, thành_tiền), tổng_tiền.\n"
                "Nếu trường nào không có thông tin, hãy để null hoặc chuỗi rỗng. Chỉ trả về chuỗi JSON thuần túy, không bọc trong ```json."
            )
            
            response = ai_gateway.process_text(text_input=raw_text, system_prompt=system_prompt)
            
            if response and response.get("status") == "success":
                content_json = response.get("data")
                project.content_json = content_json
                project.save()
                
                DataEntry.objects.update_or_create(
                    project=project,
                    name=f"OCR_Text_{project.name}",
                    defaults={
                        'refine_status': 'REFINED',
                        'category': 'OCR_DOC',
                        'processed_content': f"--- NỘI DUNG OCR GỐC ---\n{raw_text}\n\n--- CẤU TRÚC AI TRÍCH XUẤT ---\n{content_json}",
                        'description': f"Dữ liệu bóc tách từ ảnh qua EasyOCR + AI Tinh chế."
                    }
                )
                return True, {"status": "success", "data": content_json}
            else:
                return False, f"AI tinh chế cấu trúc thất bại: {response.get('message', 'Không rõ lỗi')}"
                
        except Exception as e:
            logger.error(f"Lỗi quy trình Hybrid OCR + AI: {str(e)}")
            return False, traceback.format_exc()

    def _classify_cell(self, ws, cell):
        val_str = str(cell.value).strip().lower() if cell.value is not None else ""
        if val_str.startswith('='):
            return 'LOGIC', self._get_smart_label(ws, cell)
        if any(kw in val_str for kw in self.ui_keywords):
            return 'UI', cell.value
        if len(val_str) > 100 and not cell.value:
            return 'TRASH', "Long_Text_Garbage"
        return 'DATA', self._get_smart_label(ws, cell)

    def _get_smart_label(self, ws, cell):
        try:
            if cell.column > 1:
                l_val = ws.cell(row=cell.row, column=cell.column - 1).value
                if l_val and isinstance(l_val, str): return l_val.strip()
            if cell.row > 1:
                t_val = ws.cell(row=cell.row - 1, column=cell.column).value
                if t_val and isinstance(t_val, str): return t_val.strip()
        except: pass
        return f"Cell_{cell.coordinate}"

    def _build_field_obj(self, sheet_obj, cell, field_type, label):
        val_str = str(cell.value) if cell.value is not None else ""
        return DataField(
            sheet=sheet_obj,
            cell_address=cell.coordinate,
            field_type=field_type,
            label=label[:255] if label else None,
            value=val_str if not val_str.startswith('=') else None,
            formula=val_str if val_str.startswith('=') else None,
        )

    def _create_unified_draft(self, project, sheet_obj):
        KnowledgeDraft.objects.update_or_create(
            project=project,
            sheet=sheet_obj,
            defaults={
                'term': f"Nghiệp vụ Sheet {sheet_obj.name}",
                'content': f"Bản thảo chờ AI phân tích nghiệp vụ cho sheet {sheet_obj.name}.",
                'category': 'SHEET_LOGIC',
                'status': 'PENDING',
                'origin_metadata': sheet_obj.metadata 
            }
        )